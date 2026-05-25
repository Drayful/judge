from __future__ import annotations

from datetime import date
from pathlib import Path


def main() -> None:
    # Lazy import so the script can instruct how to install dependency.
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as e:  # pragma: no cover
        raise SystemExit(
            "Missing dependency python-docx. Install with:\n"
            "  pip install python-docx\n"
            f"Original error: {e}"
        )

    out_path = Path(__file__).resolve().parent / "weekly_report_2026-05-08.docx"

    doc = Document()

    title = doc.add_paragraph()
    run = title.add_run("Отчет за неделю по проекту Judge")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph("Период: 04.05.2026–08.05.2026")
    subtitle.runs[0].font.size = Pt(11)

    doc.add_paragraph()

    def h(text: str) -> None:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(13)

    def bullets(items: list[str]) -> None:
        for it in items:
            doc.add_paragraph(it, style="List Bullet")

    h("Кратко")
    bullets(
        [
            "Подготовлен и опубликован репозиторий проекта на GitHub (ветка main).",
            "Настроен/использован Docker Compose (Laravel Sail) для локального запуска.",
            "Добавлены/использованы сервисы импорта стартового протокола, продвижения очереди и загрузки музыки с версионированием.",
        ]
    )

    h("Функциональные изменения")
    bullets(
        [
            "Добавлен флаг авто-продвижения в категории: миграция 2026_05_04_120000_add_auto_advance_to_categories_table (поле auto_advance).",
            "Добавлена связь «активный поток/категория» для турнира: миграция 2026_05_04_160000_add_active_category_id_to_tournaments_table (FK active_category_id → categories).",
            "Сервис продвижения выступлений в категории: StreamAdvanceService (performing → done, затем scheduled → performing).",
        ]
    )

    h("Интерфейс и UX")
    bullets(
        [
            "Экран судьи для планшета: resources/views/judge/tablet.blade.php (панели A/D/E/штраф, черновик оценки, защита от повторной отправки).",
            "Экран ожидания выбора потока: resources/views/judge/tablet-wait.blade.php (периодический ping и авто-обновление).",
            "Базовый layout для табло: resources/views/components/scoreboard-layout.blade.php (темная тема, подключение Vite).",
        ]
    )

    h("Импорт/музыка")
    bullets(
        [
            "Импорт стартового протокола из Excel: StartProtocolImportService (обработка «Группа/Поток», создание категорий, участниц и выступлений, учет числа кругов/видов).",
            "Загрузка музыки с историей версий: MusicTrackUploadService (version++, деактивация предыдущей активной версии, хранение в Storage).",
            "Учет дедлайна загрузки музыки категории и проверка прав пользователя на обход дедлайна.",
        ]
    )

    h("DevOps и инфраструктура")
    bullets(
        [
            "Docker Compose (compose.yaml, Laravel Sail): сервисы laravel.test, pgsql (postgres:18-alpine), redis, mailpit, minio.",
            "Проект подготовлен к локальному поднятию через docker compose up -d --build.",
        ]
    )

    h("Безопасность")
    bullets(
        [
            "Секреты не публикуются в репозиторий: .env уже игнорируется, DEMO_CREDENTIALS.txt добавлен в .gitignore.",
        ]
    )

    h("Что дальше")
    bullets(
        [
            "Добавить рабочий README проекта (как запускать через Sail, как заполнить .env, базовый сценарий тестового запуска).",
            "Стабилизировать realtime (ScoreUpdated) и сценарий обновления табло/планшетов в прод окружении.",
            "Добавить тестовые данные/seed и короткую инструкцию для судей/секретаря.",
        ]
    )

    doc.add_paragraph()
    footer = doc.add_paragraph(f"Сформировано: {date(2026, 5, 8).isoformat()}")
    footer.runs[0].font.size = Pt(9)

    doc.save(out_path)


if __name__ == "__main__":
    main()

