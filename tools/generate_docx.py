from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]


def save_done_docx(out_path: Path) -> None:
    doc = Document()
    doc.add_heading(
        "Веб‑судейство по художественной гимнастике — что уже сделано (MVP)", level=1
    )

    doc.add_heading("1) Локальная инфраструктура (Docker)", level=2)
    items = [
        "Laravel + Breeze (Blade) + Tailwind/Vite.",
        "Docker Compose (compose.yaml) поднимает: App (laravel.test), PostgreSQL 18 (pgsql), Redis (redis), Mailpit, MinIO (S3 для музыки).",
    ]
    for it in items:
        doc.add_paragraph(it, style="List Bullet")

    doc.add_heading("2) Авторизация и роли", level=2)
    for it in [
        "Роли пользователей: admin / secretary / judge / athlete.",
        "Ограничение доступа через middleware role:.",
        "Верхнее меню показывает разделы в зависимости от роли.",
    ]:
        doc.add_paragraph(it, style="List Bullet")

    doc.add_heading("3) Сущности (модель данных)", level=2)
    for it in [
        "Tournament (турнир)",
        "Category (категория)",
        "Athlete (спортсменка)",
        "Performance (выход/проход в очереди)",
        "MusicTrack (музыка)",
        "JudgeScore (оценки судей)",
    ]:
        doc.add_paragraph(it, style="List Bullet")

    doc.add_heading("4) Ключевая логика очереди", level=2)
    for it in [
        "Одна спортсменка может встречаться в очереди несколько раз (несколько Performance).",
        "Для каждого Performance можно задать свой apparatus (снаряд).",
        "Музыка привязана к конкретному Performance (на каждый выход свой трек).",
    ]:
        doc.add_paragraph(it, style="List Bullet")

    doc.add_heading("5) Экраны", level=2)
    for it in [
        "Спортсменка: /athlete/music — выбор выхода и загрузка музыки для этого выхода (S3/MinIO).",
        "Секретарь: /secretary — список категорий; /secretary/categories/{id}/queue — очередь, вызов следующей, старт/финиш, скачивание музыки по выходу.",
        "Судья: /judge — список категорий; /judge/categories/{id} — ввод D/A/E/penalty, расчёт итога.",
        "Табло: /scoreboard/categories/{id} — публичное табло; live‑обновление через /scoreboard/categories/{id}/live.",
    ]:
        doc.add_paragraph(it, style="List Bullet")

    doc.add_heading("6) Демо‑доступы", level=2)
    doc.add_paragraph(
        "Создан файл DEMO_CREDENTIALS.txt с логинами/паролями для теста.",
        style="List Bullet",
    )

    doc.add_paragraph(
        "Документ сформирован автоматически для согласования и ТЗ.",
    )

    doc.save(out_path)


def save_questions_docx(out_path: Path) -> None:
    doc = Document()
    doc.add_heading("Вопросы заказчику для отличного веб‑судейства", level=1)

    def section(title: str, questions: list[str]) -> None:
        doc.add_heading(title, level=2)
        for q in questions:
            doc.add_paragraph(q, style="List Number")

    section(
        "1) Регламент и подсчёт",
        [
            "Какая схема судейства нужна: D+E или D+A+E?",
            "Сколько судей на каждой панели (D/E/A/штраф)?",
            "Как агрегировать оценки: среднее всех / среднее без max‑min / медиана?",
            "Как считаются штрафы: кто вводит и как суммируются?",
            "Какая точность и округление: 0.05 / 0.1 / 0.01?",
            "Нужны ли пороги расхождения (флаг главному судье) и правила корректировки?",
        ],
    )
    section(
        "2) Структура стартов",
        [
            "Категория всегда на один снаряд или в одной категории могут быть разные снаряды по выходам?",
            "Нужны ли многоборье, квалификация, финалы по видам, перенос баллов?",
            "Нужны ли повторные попытки/перезапуски и как это отражать в протоколе?",
        ],
    )
    section(
        "3) Процесс в день соревнований",
        [
            "Сколько ковров/потоков одновременно?",
            "Нужны ли статусы: on deck / performing / done / published и кто ими управляет?",
            "Публикация результатов сразу или только после утверждения главным судьёй?",
            "Нужны ли апелляции/протесты: кто подаёт, сроки, как фиксировать решения?",
        ],
    )
    section(
        "4) Музыка",
        [
            "Форматы и максимальный размер файла? Разрешить ли несколько файлов на один выход?",
            "Дедлайн загрузки/замены музыки? Нужна ли история версий?",
            "Как секретарь должен использовать музыку: скачать файл / встроенный плеер / отдельный режим «пульт»?",
        ],
    )
    section(
        "5) Роли и доступы",
        [
            "Полный список ролей: админ, организатор, главный судья, секретарь, судьи по панелям, тренер, спортсменка, зритель?",
            "Что видит тренер/спортсменка: оценки сразу или только после публикации?",
            "Нужны ли ограничения «судья видит только свою панель»?",
        ],
    )
    section(
        "6) Табло и отчёты",
        [
            "Что показывать публично: только итог/место или ещё D/A/E/штраф?",
            "Нужны ли страницы расписания, участников, клубов?",
            "Нужны ли печатные формы/протоколы и экспорт в PDF/Excel?",
        ],
    )
    section(
        "7) Импорт/экспорт и интеграции",
        [
            "Нужен импорт из Excel (участники/старт‑листы)? Есть ли текущий шаблон?",
            "Нужна интеграция с существующей регистрацией/системой федерации?",
        ],
    )
    section(
        "8) Нагрузка и надёжность",
        [
            "Пиковая нагрузка: сколько одновременных судей/секретарей/загрузок?",
            "Требования к скорости обновления табло?",
            "Политика хранения музыки: сколько хранить, архив, удаление?",
            "Требования к логам и аудиту изменений оценок?",
        ],
    )

    doc.add_paragraph(
        "Важно: попросите заказчика прислать регламент/пример протокола (фото/скан/файл). Это ускорит точную реализацию подсчёта."
    )
    doc.save(out_path)


def main() -> None:
    # Use ASCII filenames to avoid Windows console encoding issues.
    save_done_docx(ROOT / "01_done.docx")
    save_questions_docx(ROOT / "02_questions.docx")
    print("ok")


if __name__ == "__main__":
    main()

