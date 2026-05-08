from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]


def h(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def num(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def add_checklist(doc: Document, steps: list[str]) -> None:
    for s in steps:
        doc.add_paragraph(s, style="List Number")


def main() -> None:
    doc = Document()

    h(doc, "Отчёт: что реализовано и как проверить (по ответам заказчика)", 1)
    doc.add_paragraph(
        "Документ сформирован автоматически. "
        "Дата: " + datetime.now().strftime("%d.%m.%Y %H:%M")
    )

    h(doc, "0) Базовая проверка, что окружение поднято", 2)
    bullet(doc, "Откройте сайт: http://localhost (Laravel Sail).")
    bullet(doc, "Mailpit: http://localhost:8025")
    bullet(doc, "MinIO console: http://localhost:9001 (логин/пароль из .env)")
    h(doc, "Как проверить", 3)
    add_checklist(
        doc,
        [
            "Запустите контейнеры: docker compose up -d",
            "Откройте http://localhost — должна открыться главная страница.",
            "Войдите одним из демо‑пользователей (см. DEMO_CREDENTIALS.txt).",
        ],
    )

    h(doc, "1) Единый светлый дизайн (Blade компоненты)", 2)
    bullet(doc, "Приведены к единому стилю: welcome, dashboard, профайл, ключевые экраны ролей.")
    bullet(doc, "Используются компоненты: x-card, x-badge, x-flash.")
    h(doc, "Как проверить", 3)
    add_checklist(
        doc,
        [
            "Откройте / (welcome) — должен быть светлый лендинг без стандартной Laravel-заглушки.",
            "Откройте /dashboard — карточки в одном стиле.",
            "Откройте /profile — блоки профиля в x-card стиле.",
        ],
    )

    h(doc, "2) Очередь и статусы выступления (секретарь)", 2)
    bullet(doc, "Очередь по категории, вызов следующей, старт/финиш.")
    bullet(doc, "Статусы: scheduled / on_deck / performing / done / approved / published (+ inquiry/under_review при апелляции).")
    h(doc, "Как проверить", 3)
    add_checklist(
        doc,
        [
            "Зайдите секретарём: /secretary → выберите категорию → очередь.",
            "Нажмите «Вызвать следующую» — у первой scheduled станет on_deck.",
            "Нажмите «Старт» — статус станет performing.",
            "Нажмите «Финиш» — статус станет done.",
        ],
    )

    h(doc, "3) Судейские панели и подсчёт D/A/E + штрафы", 2)
    bullet(doc, "Роли судей по панелям: D (DB/DA), A, E, штрафы (line/time/music).")
    bullet(doc, "Подсчёт: D = avg(DB) + avg(DA); A/E: drop high/low, avg середины; Penalties: сумма; Total = D + A + E − Penalties.")
    h(doc, "Как проверить", 3)
    add_checklist(
        doc,
        [
            "Зайдите судьёй: /judge → выберите категорию.",
            "Введите оценки своей панели (кнопка OK).",
            "Нажмите «Итог» — появится рассчитанный total.",
            "Зайдите под другой судейской ролью и внесите остальные панели — итог обновится после пересчёта.",
        ],
    )

    h(doc, "4) Workflow approve/publish (Superior Jury / Head Judge)", 2)
    bullet(doc, "Утверждение (approve) и публикация (publish) результата для табло.")
    bullet(doc, "Табло показывает только published выступления.")
    h(doc, "Как проверить", 3)
    add_checklist(
        doc,
        [
            "Зайдите Superior Jury / Head Judge: /judge → категория.",
            "После подсчёта нажмите «Утвердить» (approved_at).",
            "Нажмите «Публикация» (published_at) — результат пойдёт на табло.",
            "Откройте /scoreboard/categories/{id} — увидите опубликованные строки.",
        ],
    )

    h(doc, "5) Табло (live обновление без перезагрузки)", 2)
    bullet(doc, "Live endpoint: /scoreboard/categories/{id}/live (polling).")
    bullet(doc, "Добавлено: клуб, снаряд, статусы inquiry.")
    h(doc, "Как проверить", 3)
    add_checklist(
        doc,
        [
            "Откройте табло: /scoreboard/categories/1 (или нужный id).",
            "Сделайте publish нового выступления — строка появится на табло в течение ~2 секунд.",
        ],
    )

    h(doc, "6) Музыка по выступлению (primary + backup + история версий)", 2)
    bullet(doc, "Музыка привязана к конкретному Performance.")
    bullet(doc, "Поддержка основного и резервного файла (primary/backup), версионирование и история замен.")
    h(doc, "Как проверить", 3)
    add_checklist(
        doc,
        [
            "Зайдите спортсменкой: /athlete/music.",
            "Выберите выступление → тип «Основной» → загрузите файл.",
            "Загрузите второй раз «Основной» — старая версия станет неактивной, новая станет активной.",
            "Загрузите «Резервный» — появится ссылка на backup.",
            "В очереди секретаря появятся ссылки «Основной»/«Резерв».",
        ],
    )

    h(doc, "7) Дедлайн на замену музыки", 2)
    bullet(doc, "Добавлено поле категории music_deadline_at.")
    bullet(doc, "После дедлайна спортсменка не может менять музыку; админ может.")
    h(doc, "Как проверить", 3)
    add_checklist(
        doc,
        [
            "Выставьте categories.music_deadline_at в прошлое (через БД).",
            "Попробуйте загрузить музыку спортсменкой — получите сообщение о дедлайне.",
            "Попробуйте загрузить админом — загрузка разрешена.",
        ],
    )

    h(doc, "8) Inquiry (апелляции) submitted → under_review → decided", 2)
    bullet(doc, "Секретарь создаёт inquiry для выступления.")
    bullet(doc, "Superior Jury/Head Judge переводит в under_review и выносит решение.")
    bullet(doc, "Статус inquiry отображается в очереди, в судейском экране и на табло.")
    h(doc, "Как проверить", 3)
    add_checklist(
        doc,
        [
            "Зайдите секретарём в очередь категории и создайте inquiry (кнопка «Создать»).",
            "Зайдите Superior Jury/Head Judge в судейский экран категории.",
            "Нажмите «Under review», затем выберите решение и нажмите «Decide».",
            "Проверьте, что бейджи inquiry обновились.",
        ],
    )

    h(doc, "Техническая заметка: как запускать artisan команды", 2)
    bullet(doc, "На Windows локально php artisan migrate может не видеть хост pgsql.")
    bullet(doc, "Используйте: docker compose exec -T laravel.test php artisan <команда>")

    out_path = ROOT / "03_report.docx"
    doc.save(out_path)
    print(str(out_path))


if __name__ == "__main__":
    main()

